"""If you want use a multiple language for autoCorrect System use this code"""

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from symspellpy import SymSpell, Verbosity
from docx import Document
from PyPDF2 import PdfReader
from fpdf import FPDF
import os
import tempfile
import uuid
import re
from typing import List, Dict

# Initialize FastAPI app
app = FastAPI()

# CORS Configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ====================
# Core Spell Checking
# ====================
class EnhancedSpellChecker:
    def __init__(self):
        self.spell = SymSpell()
        dictionary_path = "frequency_dictionary_en_82_765.txt"  # Path to frequency dictionary
        bigram_path = "bigram_dictionary_en_243_342.txt"  # Path to bigram dictionary
        self.spell.load_dictionary(dictionary_path, 0, 1)
        self.spell.load_bigram_dictionary(bigram_path, 0, 2)

    def check_text(self, text: str) -> List[Dict]:
        """Check spelling and return suggestions with positions."""
        words = re.findall(r'\w+|\W+', text)
        results = []
        index = 0

        for word in words:
            start_index = index
            index += len(word)  # Track word position

            if not word.strip() or not word.isalpha():
                results.append({
                    "word": word,
                    "is_correct": True,
                    "suggestions": [],
                    "start_index": start_index
                })
            else:
                suggestions = self.spell.lookup(word, Verbosity.CLOSEST, max_edit_distance=2, include_unknown=True)
                is_correct = suggestions[0].term == word if suggestions else True
                results.append({
                    "word": word,
                    "is_correct": is_correct,
                    "suggestions": [s.term for s in suggestions[:3] if s.term != word] if not is_correct else [],
                    "start_index": start_index
                })
        return results

# Initialize spell checker
spell_checker = EnhancedSpellChecker()

# ====================
# File Processing
# ====================
class FileProcessor:
    @staticmethod
    def extract_text(file_path: str) -> str:
        """Extract text from various file formats"""
        file_ext = os.path.splitext(file_path)[1][1:].lower()

        if file_ext == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read(), file_ext
        elif file_ext == 'pdf':
            reader = PdfReader(file_path)
            return "\n".join(page.extract_text() or "" for page in reader.pages), file_ext
        elif file_ext == 'docx':
            doc = Document(file_path)
            return "\n".join(para.text for para in doc.paragraphs), file_ext
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

    @staticmethod
    def create_corrected_file(text: str, original_ext: str, output_path: str) -> str:
        """Create corrected file in original format (TXT, DOCX, PDF)."""
        if original_ext == 'txt':
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
        elif original_ext == 'docx':
            doc = Document()
            doc.add_paragraph(text)
            doc.save(output_path)
        elif original_ext == 'pdf':
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            for line in text.split("\n"):
                pdf.multi_cell(190, 10, line)
            pdf.output(output_path)
        else:
            original_ext = 'txt'
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
        return original_ext

# ====================
# API Models
# ====================
class TextRequest(BaseModel):
    text: str

# ====================
# API Endpoints
# ====================
@app.post("/check-text")
async def check_text(request: TextRequest):
    """Real-time spelling check with word positions."""
    results = spell_checker.check_text(request.text)
    return {"results": results}

@app.post("/check-file")
async def check_file(file: UploadFile = File(...)):
    """Endpoint for checking uploaded files"""
    try:
        # Validate file type
        file_ext = os.path.splitext(file.filename)[1][1:].lower()
        if file_ext not in ['txt', 'pdf', 'docx']:
            raise HTTPException(400, "Unsupported file type (only txt, pdf, docx allowed)")

        # Save to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_ext}") as tmp:
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name

        # Process file
        text, original_ext = FileProcessor.extract_text(tmp_path)
        results = spell_checker.check_text(text)

        # Generate corrected text
        corrected_text = " ".join(
            word_data['suggestions'][0] if word_data['suggestions'] else word_data['word']
            if not word_data['is_correct'] else word_data['word']
            for word_data in results
        )

        # Create corrected file
        output_filename = f"corrected_{uuid.uuid4().hex[:8]}.{original_ext}"
        output_path = os.path.join(tempfile.gettempdir(), output_filename)
        final_ext = FileProcessor.create_corrected_file(corrected_text, original_ext, output_path)

        # Cleanup and prepare response
        os.unlink(tmp_path)
        corrections_count = sum(1 for r in results if not r['is_correct'])

        return {
            "original_text": text,
            "corrected_text": corrected_text,
            "download_url": f"/download-corrected/{output_filename}",
            "file_type": final_ext,
            "corrections_count": corrections_count
        }

    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(500, f"Error processing file: {str(e)}")

@app.get("/download-corrected/{filename}")
async def download_corrected(filename: str):
    """Endpoint for downloading corrected files"""
    file_path = os.path.join(tempfile.gettempdir(), filename)
    if not os.path.exists(file_path):
        raise HTTPException(404, "File not found")

    return FileResponse(
        file_path,
        filename=f"corrected_{filename}",
        media_type="application/octet-stream"
    )